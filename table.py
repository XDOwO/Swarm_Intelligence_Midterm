
import pandas as pd
from docx import Document

# 重新讀取已建立的 CSV 檔案
func_df = pd.read_csv("ortho_win_count_by_function.csv")
dim_df = pd.read_csv("ortho_win_count_by_dimension.csv")
algo_df = pd.read_csv("ortho_win_count_by_algorithm.csv")

# 加入勝率欄位 (%)
for df in [func_df, dim_df, algo_df]:
    df["O勝率(%)"] = (df["O"] / df["總"] * 100).round(2)
    df["N勝率(%)"] = (df["N"] / df["總"] * 100).round(2)

# 建立 Word 文件
doc = Document()
doc.add_heading("正交基底成效統計報告", level=1)

def add_table_from_df(doc, title, df):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = 'Table Grid'

    # 表頭
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)

    # 表格內容
    for i, row in df.iterrows():
        for j, col in enumerate(df.columns):
            table.cell(i + 1, j).text = str(row[col])

# 寫入三張表格
add_table_from_df(doc, "依函數分類的勝負次數統計", func_df)
add_table_from_df(doc, "依維度分類的勝負次數統計", dim_df)
add_table_from_df(doc, "依演算法分類的勝負次數統計", algo_df)

# 儲存 Word 文件
doc.save("ortho_win_summary.docx")
